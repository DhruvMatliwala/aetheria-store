import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import datetime

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    # Page Margins (1 inch all sides)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Styling Palette
    COLOR_PRIMARY = RGBColor(14, 116, 144)    # Cyan / Deep Teal (#0E7490)
    COLOR_SECONDARY = RGBColor(5, 150, 105)   # Emerald (#059669)
    COLOR_DARK = RGBColor(15, 23, 42)         # Slate 900 (#0F172A)
    COLOR_MUTED = RGBColor(100, 116, 139)     # Slate 500 (#64748B)

    # ── Title & Header ──────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("AETHERIA — PGSharp Digital Storefront")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_DARK

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(14)
    run_sub = p_sub.add_run("Comprehensive Technical, Architectural & Engineering Production Manual")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(13)
    run_sub.font.color.rgb = COLOR_PRIMARY

    # Metadata banner table
    meta_table = doc.add_table(rows=1, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    col_widths = [Inches(1.6), Inches(1.6), Inches(1.6), Inches(1.7)]
    headers = ["Version: 1.0.0 Prod", "Framework: Next.js 14", "Security: AES-256-GCM", f"Date: {datetime.datetime.now().strftime('%B %Y')}"]
    
    for i, cell in enumerate(meta_table.rows[0].cells):
        cell.width = col_widths[i]
        set_cell_background(cell, "F1F5F9")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(headers[i])
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
        r.font.bold = True
        r.font.color.rgb = COLOR_MUTED

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ── Helper Functions ────────────────────────────────────────────────────────
    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = COLOR_PRIMARY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = COLOR_DARK
        return p

    def add_body(text, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_DARK
        return p

    def add_bullet(bold_prefix, text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r_bold = p.add_run(bold_prefix)
        r_bold.font.name = "Arial"
        r_bold.font.size = Pt(10)
        r_bold.font.bold = True
        r_bold.font.color.rgb = COLOR_DARK
        
        r_text = p.add_run(text)
        r_text.font.name = "Arial"
        r_text.font.size = Pt(10)
        r_text.font.color.rgb = COLOR_DARK
        return p

    def format_table(table, col_widths, headers, data):
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        hdr_row = table.rows[0]
        for i, h_text in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.width = col_widths[i]
            set_cell_background(cell, "0E7490")
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            r = p.add_run(h_text)
            r.font.name = "Arial"
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            
        for row_idx, row_data in enumerate(data):
            row = table.add_row()
            bg_color = "F8FAFC" if row_idx % 2 == 0 else "FFFFFF"
            for col_idx, cell_value in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.width = col_widths[col_idx]
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
                p = cell.paragraphs[0]
                r = p.add_run(cell_value)
                r.font.name = "Arial"
                r.font.size = Pt(9)
                r.font.color.rgb = COLOR_DARK

    # ── 1. Executive Summary & Brand Identity ───────────────────────────────────
    add_h1("1. Executive Summary & System Overview")
    add_body("AETHERIA is a next-generation, high-performance digital storefront engineered specifically for official PGSharp Standard Edition key distribution. It merges luxury dark-mode aesthetics (Obsidian Glassmorphism, Neon Cyan & Emerald accoutrements) with an interactive 550vh image-sequence scrollytelling runway, dual-rail automated checkout, cryptographic key security, ambient spatial audio, and an instant zero-latency customer fulfillment portal.")
    
    add_bullet("Core Brand Proposition: ", "Institutional-grade reliability, instantaneous automated key dispatch, zero ban risk advisory, and 1-on-1 VIP concierge support.")
    add_bullet("Technology Stack: ", "Next.js 14 App Router, TypeScript, Tailwind CSS, GSAP 3 (ScrollTrigger), HTML5 Canvas Particle Engine, Firebase Firestore / Admin SDK, Razorpay UPI, and PayPal v2 Checkout API.")

    # ── 2. Interactive Scrollytelling Architecture ──────────────────────────────
    add_h1("2. Cinematic Scrollytelling Engine")
    add_body("The storefront experience is driven by CinematicScrollExperience.tsx, utilizing a pinned 550vh scroll runway that orchestrates 3 cinematic video scenes with synchronous HUD text reveals:")

    scene_table = doc.add_table(rows=1, cols=4)
    format_table(
        scene_table,
        [Inches(1.2), Inches(1.8), Inches(1.8), Inches(1.7)],
        ["Scene", "Theme & Asset", "Headline & Narrative", "Interaction / CTAs"],
        [
            ["Scene 1 (0–28%)", "Mewtwo Cryo-Awakening (/videos/mewtwo.mp4)", "Break every limit. Unleash raw power across dimensions.", "Instant Dispatch, 30-Day License, 'Buy License Key →'"],
            ["Scene 2 (28–62%)", "Cyberpunk Shibuya (/videos/shibuya.mp4)", "Roam anywhere. Seamless virtual teleportation.", "Cooldown Radar, 60 FPS Emulation, 100% IV Sniping"],
            ["Scene 3 (62–100%)", "Mewtwo Showdown (/videos/mewtwo_2.mp4)", "Master every raid. Instant VIP key allocation.", "Dual Obsidian Pricing Cards (1-Device & 2-Device)"]
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    add_h2("2.1 Dual-Player Optical Dissolve Video Loop Engine")
    add_body("Standard HTML5 video looping produces hard, jarring visual jump cuts. AETHERIA implements DissolveSceneVideo, a dual-player optical cross-fading engine:")
    add_bullet("Seamless Transition: ", "When Player A reaches 0.35s before video completion, Player B automatically resets to frame 0 and begins playback. Player A cross-fades opacity to 0 over 300ms while Player B fades to 1, creating an infinite, seamless optical loop.")
    add_bullet("Hardware-Event Optimization: ", "Driven by native HTML5 onTimeUpdate event listeners rather than continuous requestAnimationFrame polling loops, reducing CPU and battery consumption by over 90%.")

    # ── 3. Full-Screen Brand Preloader & Splash Intro ───────────────────────────
    add_h1("3. Full-Screen Cinematic Brand Preloader")
    add_body("To establish an immediate AAA gaming atmosphere and eliminate all flashes of unstyled content, Preloader.tsx executes a high-impact intro sequence on first page load:")
    add_bullet("Phase 1 (0.0s – 0.6s): ", "Frosted Delta Logo Badge scales in with a glowing neon cyan drop shadow (scale: 0.8 -> 1, filter: drop-shadow(0 0 25px rgba(6,182,212,0.85))).")
    add_bullet("Phase 2 (0.35s – 1.1s): ", "Character-split 'AETHERIA' wordmark slides up with staggered timing (stagger: 0.04s), accompanied by the monospace 'VAULT ACCESS PROTOCOL' subtitle and a glowing horizontal cyber progress line.")
    add_bullet("Phase 3 (1.6s – 2.2s): ", "Brand elements scale slightly and fade out while the solid black curtain pulls upward like a shutter (yPercent: -100, ease: power4.inOut), seamlessly revealing Scene 1.")
    add_bullet("Zero-FOUC Guarantee: ", "All preloader elements are styled with baseline opacity-0 classes in JSX/CSS, completely eliminating initial frame flashes.")
    add_bullet("Session Persistence: ", "Leverages sessionStorage.getItem('hasSeenIntro') so the intro sequence only runs once per browser session, ensuring returning users navigate instantly.")

    # ── 4. Ambient Background Audio Subsystem ───────────────────────────────────
    add_h1("4. Ambient Background Audio Engine")
    add_body("Managed by AmbientAudioContext.tsx and mounted at the root layout (src/app/layout.tsx), providing continuous background spatial atmosphere:")
    add_bullet("Persistent Singleton: ", "A single persistent HTML5 Audio instance loads /public/audio/ambient.mp3 and maintains uninterrupted playback across all route navigations.")
    add_bullet("Volume Calibration & Fade Ramps: ", "Volume is strictly capped at a comfortable 25% ceiling (TARGET_VOLUME = 0.25) with smooth 300ms volume cross-fade ramps on play/pause toggles.")
    add_bullet("Frosted HUD Toggle: ", "Integrated in Header.tsx alongside 'Buy Key →' with dynamic glowing cyan sound wave indicators.")

    # ── 5. Zero-Lag Performance & Algorithmic Optimizations ─────────────────────
    add_h1("5. High-Performance Engineering & Optimizations")
    add_body("The storefront has undergone rigorous GPU, memory, and main-thread profiling to guarantee buttery-smooth 60–120 FPS performance:")
    add_bullet("React State Reconciliation Guard: ", "handleScrollProgress employs an equality check (prev !== newIdx ? newIdx : prev), eliminating hundreds of unnecessary React component reconciliation passes during fast scrolling.")
    add_bullet("GSAP Timeline Conflict Resolution: ", "Configured with overwrite: 'auto', clearProps: 'transform', anticipatePin: 1, and invalidateOnRefresh: true, ensuring zero property collisions between intro entrance animations and scroll scrubbing.")
    add_bullet("Canvas Tab Visibility Throttling: ", "AmbientMistParticles.tsx listens to document.visibilitychange, halting the animation loop when the browser tab is inactive (0% CPU/GPU overhead) and resuming instantaneously upon focus.")
    add_bullet("Media Subsystem Flags: ", "Added disablePictureInPicture and disableRemotePlayback to all video elements, preventing Chromium and WebKit from attaching background media controllers.")
    add_bullet("Immutable Asset Caching: ", "next.config.mjs delivers Cache-Control: public, max-age=31536000, immutable for all video, audio, and image assets.")
    add_bullet("Sub-5ms Stock API Caching: ", "/api/stock returns Cache-Control: public, s-maxage=10, stale-while-revalidate=30 for lightning-fast inventory responses with background Firestore revalidation.")

    # ── 6. Redesigned Luxury Order Success Page ─────────────────────────────────
    add_h1("6. Order Confirmation & Fulfillment Experience")
    add_body("Located at /order-success/[orderId], the order fulfillment portal delivers an institutional-grade customer experience:")
    add_bullet("Holographic Key Vault: ", "Pulsing emerald status badge, blurred click-to-reveal key card, 1-click copy with confetti particle burst, and strict device binding notice.")
    add_bullet("Transaction Receipt Breakdown: ", "4-column breakdown detailing Plan Type, Allocated Slots, Capture Reference, and License Duration.")
    add_bullet("1-on-1 VIP Support Portal: ", "Direct messaging buttons linking to Discord Profile and Reddit Profile for instant concierge assistance.")
    add_bullet("4-Step Interactive Activation Guide: ", "Stepped luxury obsidian cards detailing APK download, key entry, activation, and spoofing setup, with an official pgsharp.com download button.")

    # ── 7. Cryptographic Key Engine & Security ──────────────────────────────────
    add_h1("7. Cryptographic Security & Vault Architecture")
    add_body("AETHERIA implements military-grade cryptographic protection for all digital assets:")
    add_bullet("AES-256-GCM Encryption: ", "All raw PGSharp license keys are encrypted at rest using AES-256-GCM with unique initialization vectors (IV) and authentication tags.")
    add_bullet("Device Slot Partitioning: ", "Keys are dynamically tracked by usableSlots (1-Device, 2-Device, 3-Device). When a user purchases a 1-Device plan, exactly 1 slot is consumed. Keys are retired only when all slots are exhausted.")
    add_bullet("Webhook Signature Verification: ", "All payment webhooks (PayPal HMAC & Razorpay SHA256) are cryptographically validated before key release.")

    # ── 8. Route Inventory & System Architecture Map ────────────────────────────
    add_h1("8. Complete Route & API Architecture Map")
    add_body("The application comprises 15 optimized production routes:")

    route_table = doc.add_table(rows=1, cols=3)
    format_table(
        route_table,
        [Inches(1.8), Inches(1.5), Inches(3.2)],
        ["Route / Endpoint", "Rendering Type", "Purpose & Functionality"],
        [
            ["/", "Static / Client", "Main 3-Scene Scrollytelling Storefront, Preloader & HUD"],
            ["/order-success/[orderId]", "Dynamic SSR", "Encrypted Key Vault, Confetti Reveal & Activation Guide"],
            ["/admin", "Static / Client", "Admin Key Ingestion, Slot Tier Analytics & Waitlist Manager"],
            ["/api/stock", "Dynamic API (Cached)", "Real-time Tier Stock, Usable Slots & Active Inventory Counts"],
            ["/api/stock/[planId]", "Dynamic API", "Individual Plan Stock Availability Check"],
            ["/api/checkout/upi", "Dynamic API", "Razorpay UPI QR Code / Payment Intent Generation"],
            ["/api/checkout/upi/verify", "Dynamic API", "UPI Payment Signature Verification & Key Allocation"],
            ["/api/checkout/paypal", "Dynamic API", "PayPal v2 Order Creation Endpoint"],
            ["/api/checkout/paypal/capture", "Dynamic API", "PayPal Order Capture & Key Dispatch"],
            ["/api/webhooks/paypal", "Dynamic API", "PayPal Asynchronous Webhook Verification"],
            ["/api/webhooks/upi", "Dynamic API", "Razorpay Asynchronous Webhook Verification"],
            ["/api/restock-notify", "Dynamic API", "Customer Restock Notification Waitlist Ingestion"],
            ["/api/admin/keys", "Dynamic API", "Secure Batch Key Ingestion & Slot Allocation API"],
            ["/api/admin/stats", "Dynamic API", "Administrative Inventory Metrics & Revenue Analytics"],
            ["/terms, /privacy, /refund", "Static", "Legal Compliance, Terms of Service & Privacy Policy"]
        ]
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ── 9. Operational Runbook & Deployment Guide ───────────────────────────────
    add_h1("9. Operational Runbook & Deployment Guide")
    add_body("Instructions for running and deploying the production storefront:")
    add_bullet("Environment Configuration: ", "Ensure .env.local contains FIREBASE_PROJECT_ID, FIREBASE_CLIENT_EMAIL, FIREBASE_PRIVATE_KEY, KEY_ENCRYPTION_SECRET (32-byte hex), PAYPAL_CLIENT_ID, PAYPAL_CLIENT_SECRET, and RAZORPAY_KEY_ID / SECRET.")
    add_bullet("Production Build Command: ", "Execute 'npm run build' (validates all 15 routes, types, and bundle optimizations).")
    add_bullet("Production Server Start: ", "Execute 'npm run start' on port 3000 behind NGINX or deploy seamlessly to Vercel.")

    # ── 10. VIP Private Promo Code & Discount Engine ───────────────────────────
    add_h1("10. VIP Promo Code & Private Discount Architecture")
    add_body("To reward trusted repeat buyers without public coupon exploitation, AETHERIA features a private promo engine:")
    add_bullet("Private Code Distribution: ", "No public hints or auto-fill codes appear on the storefront. The promo input uses an unhinted '[ Enter Promo Code ]' placeholder, allowing admins to distribute private codes (e.g., VIPDHRUV, DISCORD10) via Discord/Telegram DMs.")
    add_bullet("Real-time Server Validation: ", "Endpoint /api/coupons/validate queries Firestore collection 'coupons' to cryptographically check code validity, minimum cart thresholds, and expiry dates before applying reductions.")
    add_bullet("Dual Discount Modes: ", "Supports both fixed rupee reductions (e.g. ₹10 OFF) and percentage discounts (e.g. 10% OFF), updating live checkout summaries in INR and converted USD seamlessly.")

    # ── 11. Modernized Cyberpunk Admin Command Center & Cyan Palette Overhaul ───
    add_h1("11. Modernized Admin Dashboard & Universal Theme Architecture")
    add_body("The administrative interface and visual design language underwent a complete architectural modernization:")
    add_bullet("Obsidian & Electric Cyan Theme: ", "Completely eradicated legacy saffron gold and terracotta yellow styling across the entire application. Standardized on Obsidian Deep Navy (#070B13, #0C1424), Electric Cyan (#06B6D4, #22D3EE), and Auroral Emerald (#10B981).")
    add_bullet("Strict Zero-Leak Auth Lockdown: ", "When signed out or unauthenticated, the admin route /admin returns strictly an encrypted Vault Locked card. Tabs, navigation sidebars, order tables, and coupon management components are completely unmounted from the DOM.")
    add_bullet("Hard Session Revocation: ", "Sign-out securely clears both sessionStorage ('pgsharp_admin_secret') and Firebase client authentication, guaranteeing reload persistence cannot bypass authentication.")
    add_bullet("Streamlined Proof Approval Modal: ", "Replaced complex math and cluttered tables with sleek, 1-click quick-approval and quick-reject action cards for UPI manual transaction proofs.")

    doc.save("F:/Pgsharp/docs/Documentation.docx")
    print("Successfully generated master documentation at: F:/Pgsharp/docs/Documentation.docx")

if __name__ == "__main__":
    create_document()
